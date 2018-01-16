#! python3
################################################################################
#
# Author: Award Solutions, Inc.
# Course: Internal Tool
# FIle: proc_createOL.py
#
################################################################################
#
# This is a simple python file takes an input file from the command line or the
# 'all' option.  The input is the name of a docx course outline (legacy), and in
# the case all is entered the program will look for all docx files in the current
# directory.  The program then copies the xlsm template to match the name given
# and populates the course-outline.xlsm file with the information obtained from
# the docx word outline.  Based on Version 2.9 template.
#
################################################################################
import csv
import docx
from docx import Document
import json
import getopt
import re
import openpyxl
from openpyxl.utils import column_index_from_string
from openpyxl import Workbook
import sys
import os
from shutil import copyfile
import time


################################################################################
#
# Function Name: getHandles
# Input: Input File name
# Output: file handles for word document, and xlsm course plan
#
################################################################################
def getHandles(inputfile, temp, inputTab, apiTab):
	#First make a copy of the template named to match the word outline
	destination = inputfile.replace("docx", "xlsm")
	#print("The template was =", temp)
	#print("The destination was=", destination)
	copyfile(temp, destination)
	#Next attempt to create a DocHandle for the current Word input file
	try:
		inputDocHandle = docx.Document(inputfile)
		coursePlan = openpyxl.load_workbook(destination, read_only=False, keep_vba=True)
		outlineWS = coursePlan.get_sheet_by_name(inputTab)
		outlineAPI = coursePlan.get_sheet_by_name(apiTab)
	except IOError:
		print("++-->> File open for: ", inputfile, " FAILED!")
		return(-1)
	return inputDocHandle, coursePlan, outlineWS, outlineAPI, destination
	#End of getHandles function


################################################################################
#
# Function Name: readPrintWordFile
# Input: Input File name
#
################################################################################
def readPrintWordFile(inputDocHandle):
	# expecting that the DOCX file has already been opened
	for paraCount in range(0, len(inputDocHandle.paragraphs)):
		print("")
		print("++-->> PARA NUMBER: ", paraCount)
		print(inputDocHandle.paragraphs[paraCount].style)
		print(inputDocHandle.paragraphs[paraCount].text)
	#inputDocHandle.save("NEW_"+fileName)
	return
	# END Of Function

################################################################################
#
# Function Name: getText
# Input: Input File name 
#
################################################################################
def getText(inputDocxFile):
	fullText = []
	paraStyle = []
	for para in inputDocxFile.paragraphs:
		fullText.append(para.text)
		paraStyle.append(para.style)
	#print("Full Text =", fullText)
	#print("length of full text =", len(fullText))
	#print("length of paraStyle =", len(paraStyle))
	#print("paraStyle =", paraStyle)
	return fullText, paraStyle
	# END Of Function


################################################################################
#
# Function Name: writeCP
# Input: Input File name, tabs, and raw data 
# Output: This populates the copied XLSM template (according to v2.1 field defs)
#
################################################################################
def writeCP(rawText, control, currentCP, outlineWS, APITab, destination):
	#The Course Name resides in 1-Budgeting 'C6' and is fixed in Word Doc
	APITab.cell(row=6, column=3).value = rawText[0]
	
	#The type, duration, and course number are next but must be parsed
	header = list(rawText[1].split('|'))
	APITab.cell(row=22, column=3).value = header[0]  #type of delivery (ILT)
	duration = list(header[1].split(' '))
	APITab.cell(row=20, column=3).value = duration[2] #duration numerical portion
	APITab.cell(row=21, column=3).value = duration[3] #duration units (days, hours, etc.)
	APITab.cell(row=7, column=3).value = header [2]  #course number
	
	#At this point the paragraphs become variable, and the next field
	#is the course description.
	paraNum = 2
	desc = []
	while 'Course Description' in str(control[paraNum]):
		desc.append(rawText[paraNum])
		paraNum += 1
	outlineWS.cell(row=25, column=2).value = ''.join(str(line) for line in desc)
	
	#Now keep walking through the rawText list, until we get to Intended Audience
	while rawText[paraNum] != 'Intended Audience':
		paraNum += 1
	audience = []
	#Now that the Intended Audience is found, collect lines until Learning Objectives
	while rawText[paraNum] != 'Learning Objectives':
		paraNum += 1
		audience.append(rawText[paraNum])
	#After hitting Learning Objectives the audience list is populated
	outlineWS.cell(row=34, column=2).value = ''.join(str(aud) for aud in audience)
	
	#Next keep iterating until the change in paragraph type is encountered
	while 'Basic Paragraph' in str(control[paraNum]):
		paraNum += 1
	k = 0
	while 'Learning Objectives' in str(control[paraNum]):
		outlineWS.cell(row=(42 + k), column=3).value = rawText[paraNum]
		k += 1
		paraNum += 1
	#Now that the Learning Objectives are finished paraNum is on the Suggested Prerequisites
	paraNum += 1
	k = 0
	while 'Learning Objectives' in str(control[paraNum]):
		outlineWS.cell(row=(103 + k), column=4).value = rawText[paraNum]
		k += 1
		paraNum += 1
	#Now that the Prerequisites are complete we get to Required Equipment
	paraNum += 1
	equip = []
	while 'Outline Heading' not in str(control[paraNum]):
		if 'Course Outline' not in rawText[paraNum]:
			print('Required Equipment =', rawText[paraNum])
			equip.append(rawText[paraNum])
		paraNum += 1
	outlineWS.cell(row=38, column=2).value = ''.join(str(e) for e in equip)
	
	#Now done with required equipment paraNum is at start of outline
	
	startCount = paraNum
	lineCount = 52
	headCount = 0
	endCount = len(rawText)
	for paraNum in range(startCount, endCount):
		if 'Heading' in str(control[paraNum]):
			lineCount = 52 + (headCount * 6)
			outlineWS.cell(row=lineCount, column=3).value = rawText[paraNum]
			headCount += 1
		elif 'Body' in str(control[paraNum]):
			lineCount += 1
			if 'Exercise' not in str(rawText[paraNum]):
				outlineWS.cell(row=lineCount, column=4).value = rawText[paraNum]
		else:
			print('There was an issue with the outline at line=', lineCount)
	print("Destination file name =", destination)
	currentCP.save(destination)
	return
	# END Of Function
	
	
################################################################################
#
# Function Name: main program
# Input: filename from the user
# Output: provide user basic guidance on providing file name
#
################################################################################
def main(argv):
	# First, let's start with validating the number of input arguments to Python
	# is more than the minimum expected.
	# The list of command line arguments passed to a Python script is in sys.argv.
	if len(sys.argv) < 2:
		print ("python proc_createOL.py <input: OL-file-name> -or- <input: All>")
		sys.exit(-1)
	
	# Just for good measure, let's output the user inputs.
	print("++-->> Input File Name:", sys.argv[1])
	
	###############################################################################
	######Here is where we branch 'All' or go to the user's file
	#Template is the current name of the Course Outline template 
	template = 'Copy of TEMPLATE_CP_COURSEID_COURSETITLE_nday.xlsm'
	OLTab = '1-Outline'
	APTab = '1-Budgeting'
	
	if sys.argv[1] == 'all':  #case where a batch process of a directory is desired
		inputFiles = []
		inputFiles = os.listdir(".")
		for file in range(0, len(inputFiles)):
			if inputFiles[file].endswith('.docx'):  #case where a file in the directory is the correct Word Outline
				rawText = []
				control = []
				#create file handles for the current word, and create a current course plan
				currentFile, currentCP, outlineWS, APITab, dest = getHandles(inputFiles[file], template, OLTab, APTab)
				
				#Next get text from the inputFile and control information
				rawText, control = getText(currentFile)
				
				#Now, take the raw text fields and control information and populate the course plan
				writeCP(rawText, control, currentCP, outlineWS, APITab, dest)
				print("Completed file ", inputFiles[file])
				

	else:	                        #case where the user input a specific filename
		rawText = []
		control = []
		#create file handles for the current word, and create a current course plan
		currentFile, currentCP, outlineWS, APITab, dest = getHandles(sys.argv[1], template, OLTab, APTab)
				
		#Next get text from the inputFile and control information
		rawText, control = getText(currentFile)
				
		#Now, take the raw text fields and control information and populate the course plan
		writeCP(rawText, control, currentCP, outlineWS, APITab, dest)
		print("Completed file ", sys.argv[1])
	# Dump the paragraph and styles of the entire DOCX file
	#print("++-START->>")
	#readPrintWordFile(currentFile)
	
# END of function

if __name__ == "__main__":
	main(sys.argv[1:])

################################################################################
#EOF